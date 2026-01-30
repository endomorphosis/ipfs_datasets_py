"""
Enhanced text extractors for native file converter.

Provides format-specific text extraction with graceful fallbacks
when optional dependencies are not available.

Phase 2 Feature: Native text extraction from common document formats.
"""

from pathlib import Path
from typing import Union, Dict, Any, Optional
import logging
from abc import ABC, abstractmethod

logger = logging.getLogger(__name__)


class ExtractionResult:
    """
    Result of text extraction.
    
    Attributes:
        text: Extracted text content
        metadata: Format-specific metadata
        success: Whether extraction succeeded
        error: Error message if failed
    """
    
    def __init__(
        self,
        text: str = "",
        metadata: Optional[Dict[str, Any]] = None,
        success: bool = True,
        error: Optional[str] = None
    ):
        self.text = text
        self.metadata = metadata or {}
        self.success = success
        self.error = error


class TextExtractor(ABC):
    """Base class for text extractors."""
    
    @abstractmethod
    def can_extract(self, file_path: Path) -> bool:
        """Check if this extractor can handle the file."""
        pass
    
    @abstractmethod
    def extract(self, file_path: Path) -> ExtractionResult:
        """Extract text from file."""
        pass
    
    @property
    @abstractmethod
    def supported_formats(self) -> list:
        """Return list of supported MIME types."""
        pass


class PDFExtractor(TextExtractor):
    """
    Extract text from PDF files.
    
    Uses pdfplumber as primary, PyPDF2 as fallback.
    """
    
    def __init__(self):
        self.pdfplumber_available = False
        self.pypdf2_available = False
        
        try:
            import pdfplumber
            self.pdfplumber_available = True
            logger.debug("pdfplumber available for PDF extraction")
        except ImportError:
            logger.debug("pdfplumber not available")
        
        try:
            import PyPDF2
            self.pypdf2_available = True
            logger.debug("PyPDF2 available as PDF fallback")
        except ImportError:
            logger.debug("PyPDF2 not available")
    
    @property
    def supported_formats(self) -> list:
        return ['application/pdf']
    
    def can_extract(self, file_path: Path) -> bool:
        """Check if we can extract from this PDF."""
        return (self.pdfplumber_available or self.pypdf2_available) and \
               file_path.suffix.lower() == '.pdf'
    
    def extract(self, file_path: Path) -> ExtractionResult:
        """Extract text from PDF."""
        # Try pdfplumber first (better quality)
        if self.pdfplumber_available:
            result = self._extract_with_pdfplumber(file_path)
            if result.success:
                return result
        
        # Fall back to PyPDF2
        if self.pypdf2_available:
            result = self._extract_with_pypdf2(file_path)
            if result.success:
                return result
        
        return ExtractionResult(
            success=False,
            error="No PDF extraction libraries available (install pdfplumber or PyPDF2)"
        )
    
    def _extract_with_pdfplumber(self, file_path: Path) -> ExtractionResult:
        """Extract using pdfplumber."""
        try:
            import pdfplumber
            
            text_parts = []
            metadata = {'method': 'pdfplumber', 'pages': 0}
            
            with pdfplumber.open(file_path) as pdf:
                metadata['pages'] = len(pdf.pages)
                metadata['pdf_info'] = pdf.metadata
                
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            text = '\n\n'.join(text_parts)
            
            return ExtractionResult(
                text=text,
                metadata=metadata,
                success=True
            )
            
        except Exception as e:
            logger.debug(f"pdfplumber extraction failed: {e}")
            return ExtractionResult(
                success=False,
                error=f"pdfplumber extraction failed: {str(e)}"
            )
    
    def _extract_with_pypdf2(self, file_path: Path) -> ExtractionResult:
        """Extract using PyPDF2."""
        try:
            import PyPDF2
            
            text_parts = []
            metadata = {'method': 'PyPDF2', 'pages': 0}
            
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                metadata['pages'] = len(pdf_reader.pages)
                metadata['pdf_info'] = pdf_reader.metadata
                
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            text = '\n\n'.join(text_parts)
            
            return ExtractionResult(
                text=text,
                metadata=metadata,
                success=True
            )
            
        except Exception as e:
            logger.debug(f"PyPDF2 extraction failed: {e}")
            return ExtractionResult(
                success=False,
                error=f"PyPDF2 extraction failed: {str(e)}"
            )


class DOCXExtractor(TextExtractor):
    """
    Extract text from DOCX files.
    
    Uses python-docx library.
    """
    
    def __init__(self):
        self.docx_available = False
        
        try:
            import docx
            self.docx_available = True
            logger.debug("python-docx available for DOCX extraction")
        except ImportError:
            logger.debug("python-docx not available")
    
    @property
    def supported_formats(self) -> list:
        return ['application/vnd.openxmlformats-officedocument.wordprocessingml.document']
    
    def can_extract(self, file_path: Path) -> bool:
        """Check if we can extract from this DOCX."""
        return self.docx_available and file_path.suffix.lower() == '.docx'
    
    def extract(self, file_path: Path) -> ExtractionResult:
        """Extract text from DOCX."""
        if not self.docx_available:
            return ExtractionResult(
                success=False,
                error="python-docx not available (install with: pip install python-docx)"
            )
        
        try:
            import docx
            
            doc = docx.Document(file_path)
            
            # Extract paragraphs
            text_parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text]
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            text = '\n\n'.join(text_parts)
            
            metadata = {
                'method': 'python-docx',
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
            }
            
            # Extract core properties if available
            if hasattr(doc, 'core_properties'):
                props = doc.core_properties
                metadata['author'] = props.author
                metadata['title'] = props.title
                metadata['created'] = str(props.created) if props.created else None
                metadata['modified'] = str(props.modified) if props.modified else None
            
            return ExtractionResult(
                text=text,
                metadata=metadata,
                success=True
            )
            
        except Exception as e:
            logger.debug(f"DOCX extraction failed: {e}")
            return ExtractionResult(
                success=False,
                error=f"DOCX extraction failed: {str(e)}"
            )


class XLSXExtractor(TextExtractor):
    """
    Extract text from XLSX files.
    
    Uses openpyxl library.
    """
    
    def __init__(self):
        self.openpyxl_available = False
        
        try:
            import openpyxl
            self.openpyxl_available = True
            logger.debug("openpyxl available for XLSX extraction")
        except ImportError:
            logger.debug("openpyxl not available")
    
    @property
    def supported_formats(self) -> list:
        return ['application/vnd.openxmlformats-officedocument.spreadsheetml.sheet']
    
    def can_extract(self, file_path: Path) -> bool:
        """Check if we can extract from this XLSX."""
        return self.openpyxl_available and file_path.suffix.lower() == '.xlsx'
    
    def extract(self, file_path: Path) -> ExtractionResult:
        """Extract text from XLSX."""
        if not self.openpyxl_available:
            return ExtractionResult(
                success=False,
                error="openpyxl not available (install with: pip install openpyxl)"
            )
        
        try:
            import openpyxl
            
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            
            text_parts = []
            total_rows = 0
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                # Add sheet name as header
                text_parts.append(f"=== Sheet: {sheet_name} ===")
                
                # Extract rows
                for row in sheet.iter_rows(values_only=True):
                    # Filter out empty rows
                    row_values = [str(cell) if cell is not None else '' for cell in row]
                    if any(val.strip() for val in row_values):
                        text_parts.append(' | '.join(row_values))
                        total_rows += 1
                
                text_parts.append("")  # Empty line between sheets
            
            text = '\n'.join(text_parts)
            
            metadata = {
                'method': 'openpyxl',
                'sheets': len(workbook.sheetnames),
                'sheet_names': workbook.sheetnames,
                'total_rows': total_rows,
            }
            
            return ExtractionResult(
                text=text,
                metadata=metadata,
                success=True
            )
            
        except Exception as e:
            logger.debug(f"XLSX extraction failed: {e}")
            return ExtractionResult(
                success=False,
                error=f"XLSX extraction failed: {str(e)}"
            )


class HTMLExtractor(TextExtractor):
    """
    Enhanced HTML text extraction.
    
    Uses BeautifulSoup for better parsing than basic approach.
    """
    
    def __init__(self):
        self.bs4_available = False
        
        try:
            from bs4 import BeautifulSoup
            self.bs4_available = True
            logger.debug("BeautifulSoup available for HTML extraction")
        except ImportError:
            logger.debug("BeautifulSoup not available")
    
    @property
    def supported_formats(self) -> list:
        return ['text/html', 'application/xhtml+xml']
    
    def can_extract(self, file_path: Path) -> bool:
        """Check if we can extract from this HTML."""
        return file_path.suffix.lower() in ['.html', '.htm', '.xhtml']
    
    def extract(self, file_path: Path) -> ExtractionResult:
        """Extract text from HTML."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                html_content = f.read()
            
            if self.bs4_available:
                return self._extract_with_beautifulsoup(html_content, file_path)
            else:
                return self._extract_basic(html_content)
            
        except Exception as e:
            logger.debug(f"HTML extraction failed: {e}")
            return ExtractionResult(
                success=False,
                error=f"HTML extraction failed: {str(e)}"
            )
    
    def _extract_with_beautifulsoup(self, html_content: str, file_path: Path) -> ExtractionResult:
        """Extract using BeautifulSoup."""
        try:
            from bs4 import BeautifulSoup
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style tags
            for script in soup(['script', 'style']):
                script.decompose()
            
            # Extract text
            text = soup.get_text(separator='\n', strip=True)
            
            # Extract metadata
            metadata = {
                'method': 'BeautifulSoup',
                'title': soup.title.string if soup.title else None,
            }
            
            # Extract meta tags
            meta_tags = {}
            for meta in soup.find_all('meta'):
                name = meta.get('name') or meta.get('property')
                content = meta.get('content')
                if name and content:
                    meta_tags[name] = content
            
            if meta_tags:
                metadata['meta_tags'] = meta_tags
            
            return ExtractionResult(
                text=text,
                metadata=metadata,
                success=True
            )
            
        except Exception as e:
            logger.debug(f"BeautifulSoup extraction failed: {e}")
            return ExtractionResult(
                success=False,
                error=f"BeautifulSoup extraction failed: {str(e)}"
            )
    
    def _extract_basic(self, html_content: str) -> ExtractionResult:
        """Basic HTML extraction without BeautifulSoup."""
        import re
        
        # Remove script and style tags
        text = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL | re.IGNORECASE)
        
        # Remove HTML tags
        text = re.sub(r'<[^>]+>', '', text)
        
        # Decode HTML entities
        import html
        text = html.unescape(text)
        
        # Clean up whitespace
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(line for line in lines if line)
        
        metadata = {'method': 'basic_regex'}
        
        return ExtractionResult(
            text=text,
            metadata=metadata,
            success=True
        )


class OfficeFormatExtractor(TextExtractor):
    """
    Extract text from additional office formats (PPT, XLS, RTF, EPUB, OD*).
    
    Uses office_format_extractors module with graceful fallbacks.
    """
    
    def __init__(self):
        # Lazy import to avoid circular dependency
        from .office_format_extractors import (
            extract_office_format,
            get_supported_office_formats,
            is_office_format_supported
        )
        self._extract = extract_office_format
        self._is_supported = is_office_format_supported
        self._formats = get_supported_office_formats()
    
    @property
    def supported_formats(self) -> list:
        return [
            'application/vnd.ms-powerpoint',  # PPT
            'application/vnd.openxmlformats-officedocument.presentationml.presentation',  # PPTX
            'application/vnd.ms-excel',  # XLS
            'application/rtf',  # RTF
            'application/epub+zip',  # EPUB
            'application/vnd.oasis.opendocument.text',  # ODT
            'application/vnd.oasis.opendocument.spreadsheet',  # ODS
            'application/vnd.oasis.opendocument.presentation',  # ODP
        ]
    
    def can_extract(self, file_path: Path) -> bool:
        """Check if we can extract from this file."""
        return self._is_supported(file_path)
    
    def extract(self, file_path: Path) -> ExtractionResult:
        """Extract text from office format file."""
        office_result = self._extract(file_path)
        
        # Convert to ExtractionResult
        return ExtractionResult(
            text=office_result.text,
            metadata=office_result.metadata,
            success=office_result.success,
            error=office_result.error
        )


class ImageExtractor(TextExtractor):
    """
    Extractor for image files with OCR support.
    
    Supports: JPEG, PNG, GIF, WebP, SVG, BMP, TIFF, ICO, APNG
    Optional: Tesseract OCR for text extraction
    """
    
    supported_formats = {
        '.jpg', '.jpeg', '.png', '.gif', '.webp',
        '.svg', '.bmp', '.tiff', '.tif', '.ico'
    }
    
    def can_extract(self, file_path: Path) -> bool:
        """Check if file is a supported image format."""
        return file_path.suffix.lower() in self.supported_formats
    
    def extract(self, file_path: Path) -> ExtractionResult:
        """Extract text from image using OCR."""
        try:
            from .image_processor import ImageProcessor
            
            processor = ImageProcessor(ocr_enabled=True)
            result = processor.extract_text(str(file_path))
            
            return ExtractionResult(
                text=result.get('text', ''),
                metadata=result.get('metadata', {}),
                success=result.get('success', False),
                error=result.get('error')
            )
            
        except ImportError:
            logger.warning("image_processor not available")
            return ExtractionResult(
                success=False,
                error="Image processing not available"
            )
        except Exception as e:
            logger.error(f"Image extraction failed: {e}")
            return ExtractionResult(
                success=False,
                error=str(e)
            )


class ExtractorRegistry:
    """
    Registry of available text extractors.
    
    Manages format-specific extractors and routing.
    """
    
    def __init__(self):
        self.extractors = []
        self._register_default_extractors()
    
    def _register_default_extractors(self):
        """Register all default extractors."""
        self.extractors = [
            PDFExtractor(),
            DOCXExtractor(),
            XLSXExtractor(),
            HTMLExtractor(),
            OfficeFormatExtractor(),  # Additional office formats
            ImageExtractor(),  # Image formats with OCR
        ]
    
    def get_extractor(self, file_path: Union[str, Path]) -> Optional[TextExtractor]:
        """
        Get appropriate extractor for file.
        
        Args:
            file_path: Path to file
            
        Returns:
            TextExtractor instance or None
        """
        file_path = Path(file_path)
        
        for extractor in self.extractors:
            if extractor.can_extract(file_path):
                return extractor
        
        return None
    
    def extract(self, file_path: Union[str, Path]) -> ExtractionResult:
        """
        Extract text from file using appropriate extractor.
        
        Args:
            file_path: Path to file
            
        Returns:
            ExtractionResult
        """
        file_path = Path(file_path)
        
        extractor = self.get_extractor(file_path)
        
        if extractor is None:
            return ExtractionResult(
                success=False,
                error=f"No extractor available for {file_path.suffix}"
            )
        
        return extractor.extract(file_path)
    
    def get_supported_formats(self) -> list:
        """Get all supported formats across all extractors."""
        formats = set()
        for extractor in self.extractors:
            formats.update(extractor.supported_formats)
        return sorted(formats)


# Singleton registry
_default_registry = None

def get_registry() -> ExtractorRegistry:
    """Get singleton extractor registry."""
    global _default_registry
    if _default_registry is None:
        _default_registry = ExtractorRegistry()
    return _default_registry


def extract_text(file_path: Union[str, Path]) -> ExtractionResult:
    """
    Convenience function to extract text from file.
    
    Args:
        file_path: Path to file
        
    Returns:
        ExtractionResult
        
    Example:
        result = extract_text('document.pdf')
        if result.success:
            print(result.text)
            print(f"Pages: {result.metadata.get('pages')}")
    """
    return get_registry().extract(file_path)

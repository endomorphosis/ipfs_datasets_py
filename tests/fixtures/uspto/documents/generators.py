"""
Compact synthetic document generators for PATLAW-031 extraction tests.

Produces minimal PDF/DOCX/archive bytes that exercise:
  - native-text pages with filing metadata
  - image-only / scanned pages (OCR injection path)
  - DOCX structure (paragraphs, tables, core properties)
  - password-protected, corrupt, oversize, and plain archive bounds

Canaries are synthetic markers — not real confidential filings.
Prefer generators over bulk golden dumps.
"""

from __future__ import annotations

import hashlib
import io
import zipfile
from pathlib import Path
from typing import Any, Dict, Optional, Union

try:
    import fitz  # PyMuPDF
except ImportError:  # pragma: no cover
    fitz = None

try:
    from PIL import Image, ImageDraw
except ImportError:  # pragma: no cover
    Image = None
    ImageDraw = None

try:
    from docx import Document as DocxDocument
except ImportError:  # pragma: no cover
    DocxDocument = None

NATIVE_CANARY = "SYNTHETIC-NATIVE-EXTRACTION-ABSTRACT-PARAGRAPH"
SCANNED_CANARY = "SYNTHETIC-SCANNED-EXTRACTION-REQ-112"
RECEIPT_CANARY = "SYNTHETIC-ACK-RECEIPT-ID-A1B2C3D4"
DOCX_CANARY = "SYNTHETIC-DOCX-CLAIM-1-PREAMBLE"
FORM_CANARY = "SYNTHETIC-FEE-TABLE-FORM-SB08"


def sha256_hex(data: bytes | str) -> str:
    if isinstance(data, str):
        data = data.encode("utf-8")
    return hashlib.sha256(data).hexdigest()


def _require_fitz() -> None:
    if fitz is None:
        raise RuntimeError("PyMuPDF (fitz) required for PDF fixtures")


def _require_pil() -> None:
    if Image is None:
        raise RuntimeError("Pillow required for scanned PDF fixtures")


def _require_docx() -> None:
    if DocxDocument is None:
        raise RuntimeError("python-docx required for DOCX fixtures")


def _text_image(text: str, *, width: int = 420, height: int = 560) -> bytes:
    _require_pil()
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    y = 40
    for line in text.split("\n"):
        draw.text((20, y), line, fill="black")
        y += 28
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=55, optimize=True)
    return buf.getvalue()


def build_native_pdf_with_metadata(
    path: Union[str, Path, None] = None,
    *,
    text: str = NATIVE_CANARY,
    application_number: str = "16/123,456",
    form_number: str = "PTO/SB/08",
    title: str = "Synthetic Office Action Excerpt",
) -> bytes:
    """Single-page native-text PDF with metadata and filing cues."""
    _require_fitz()
    doc = fitz.open()
    try:
        page = doc.new_page(width=612, height=792)
        page.insert_text((72, 72), "UNITED STATES PATENT AND TRADEMARK OFFICE", fontsize=11)
        page.insert_text((72, 100), f"Application Number: {application_number}", fontsize=11)
        page.insert_text((72, 120), f"Form {form_number}", fontsize=11)
        page.insert_text((72, 140), text, fontsize=12)
        page.insert_text((72, 180), f"Acknowledgement Receipt ID: {RECEIPT_CANARY}", fontsize=11)
        page.insert_text((72, 210), FORM_CANARY, fontsize=10)
        # Simple two-column fee table as text (table detector may or may not fire).
        page.insert_text((72, 250), "Fee Item | Amount", fontsize=10)
        page.insert_text((72, 270), "Basic Filing | 320", fontsize=10)
        page.insert_text((72, 290), "Search | 700", fontsize=10)
        doc.set_metadata(
            {
                "title": title,
                "author": "synthetic-fixture",
                "subject": "PATLAW-031",
                "creator": "uspto-documents-generators",
            }
        )
        pdf_bytes = doc.tobytes()
    finally:
        doc.close()
    if path is not None:
        Path(path).parent.mkdir(parents=True, exist_ok=True)
        Path(path).write_bytes(pdf_bytes)
    return pdf_bytes


def build_scanned_image_only_pdf(
    path: Union[str, Path, None] = None,
    *,
    text: str = SCANNED_CANARY,
) -> bytes:
    """Image-only PDF page (no selectable text layer)."""
    _require_fitz()
    _require_pil()
    img_bytes = _text_image(text)
    doc = fitz.open()
    try:
        page = doc.new_page(width=420, height=560)
        page.insert_image(page.rect, stream=img_bytes)
        pdf_bytes = doc.tobytes()
    finally:
        doc.close()
    if path is not None:
        Path(path).parent.mkdir(parents=True, exist_ok=True)
        Path(path).write_bytes(pdf_bytes)
    return pdf_bytes


def build_password_pdf(
    path: Union[str, Path, None] = None,
    *,
    text: str = "SYNTHETIC-PASSWORD-PROTECTED",
    user_password: str = "secret",
) -> bytes:
    """Password-protected PDF (empty password must not open)."""
    _require_fitz()
    doc = fitz.open()
    try:
        page = doc.new_page()
        page.insert_text((72, 72), text, fontsize=12)
        # owner + user password; permissions empty
        pdf_bytes = doc.tobytes(
            encryption=fitz.PDF_ENCRYPT_AES_256,
            user_pw=user_password,
            owner_pw=user_password + "-owner",
        )
    finally:
        doc.close()
    if path is not None:
        Path(path).parent.mkdir(parents=True, exist_ok=True)
        Path(path).write_bytes(pdf_bytes)
    return pdf_bytes


def build_corrupt_pdf() -> bytes:
    """Truncated / corrupt PDF-like bytes."""
    return b"%PDF-1.4\n1 0 obj<<>>endobj\n%%CORRUPT_TRUNCATED"


def build_oversize_bytes(size: int = 64) -> bytes:
    """Return non-PDF payload of exact *size* (caller chooses oversize threshold)."""
    return b"X" * size


def build_docx_application(
    path: Union[str, Path, None] = None,
    *,
    claim_text: str = DOCX_CANARY,
    title: str = "Synthetic Patent Application",
    include_table: bool = True,
    include_equation_marker: bool = False,
) -> bytes:
    """Minimal DOCX with core properties, paragraphs, optional table."""
    _require_docx()
    doc = DocxDocument()
    doc.core_properties.title = title
    doc.core_properties.author = "synthetic-filer"
    doc.core_properties.subject = "PATLAW-031"
    doc.add_paragraph("UNITED STATES PATENT AND TRADEMARK OFFICE")
    doc.add_paragraph("Application Number: 16/123,456")
    doc.add_paragraph(claim_text)
    doc.add_paragraph(
        "1. A synthetic apparatus comprising a processor configured to run tests."
    )
    if include_table:
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Claim"
        table.cell(0, 1).text = "Status"
        table.cell(1, 0).text = "1"
        table.cell(1, 1).text = "Original"
    if include_equation_marker:
        # OMML not easy via python-docx; add a textual marker for difference tests.
        doc.add_paragraph("EQUATION_PLACEHOLDER: E=mc^2")
    buf = io.BytesIO()
    doc.save(buf)
    data = buf.getvalue()
    if path is not None:
        Path(path).parent.mkdir(parents=True, exist_ok=True)
        Path(path).write_bytes(data)
    return data


def build_plain_archive(
    path: Union[str, Path, None] = None,
    *,
    members: Optional[Dict[str, bytes]] = None,
) -> bytes:
    """Non-DOCX ZIP archive for bounded inventory tests."""
    members = members or {
        "readme.txt": b"synthetic archive member",
        "nested/data.bin": b"\x00\x01\x02\x03",
    }
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for name, content in members.items():
            zf.writestr(name, content)
    data = buf.getvalue()
    if path is not None:
        Path(path).parent.mkdir(parents=True, exist_ok=True)
        Path(path).write_bytes(data)
    return data


def build_zip_bomb_like(
    *,
    member_count: int = 10,
    member_size: int = 1024,
) -> bytes:
    """Many small members (for member-count bound tests, not a real zip bomb)."""
    members = {f"m{i:04d}.txt": (b"A" * member_size) for i in range(member_count)}
    return build_plain_archive(members=members)


def fixture_manifest(out_dir: Union[str, Path]) -> Dict[str, Any]:
    """Write compact fixtures to *out_dir* and return a manifest (generators, not bulk goldens)."""
    out = Path(out_dir)
    out.mkdir(parents=True, exist_ok=True)
    native = build_native_pdf_with_metadata(out / "native_metadata.pdf")
    scanned = build_scanned_image_only_pdf(out / "scanned_image_only.pdf")
    docx = build_docx_application(out / "application.docx")
    archive = build_plain_archive(out / "bundle.zip")
    corrupt = build_corrupt_pdf()
    (out / "corrupt.pdf").write_bytes(corrupt)
    return {
        "schema": "uspto.document-extraction-fixture-manifest.v1",
        "files": {
            "native_metadata.pdf": sha256_hex(native),
            "scanned_image_only.pdf": sha256_hex(scanned),
            "application.docx": sha256_hex(docx),
            "bundle.zip": sha256_hex(archive),
            "corrupt.pdf": sha256_hex(corrupt),
        },
        "canaries": {
            "native": NATIVE_CANARY,
            "scanned": SCANNED_CANARY,
            "receipt": RECEIPT_CANARY,
            "docx": DOCX_CANARY,
            "form": FORM_CANARY,
        },
    }
